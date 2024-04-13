import os
import win32com.client
from docx import Document
import PyPDF2
import re
import openpyxl
from rtf_converter import rtf_to_txt
import pytesseract
from PIL import Image
import fitz
from io import BytesIO


pytesseract.pytesseract.tesseract_cmd = r'index/Tesseract/tesseract.exe'

def doc2txt(doc_file):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    print(script_dir)
    script_dir = os.path.dirname(script_dir)
    print(script_dir)
    full_path = os.path.join(script_dir, doc_file)
    print(full_path)

    word = win32com.client.Dispatch("Word.Application")
    doc = word.Documents.Open(full_path)
    full_text = "\n".join([paragraph.Range.Text for paragraph in doc.Paragraphs])
    doc.Close()
    word.Quit()
    return full_text


def docx2txt(docx_file):
    doc = Document(docx_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text


def xlsx2txt(xlsx_file):
    wb = openpyxl.load_workbook(xlsx_file)
    text = ""
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        for row in sheet.iter_rows(values_only=True):
            row_text = '\t'.join(str(cell) for cell in row if cell is not None)
            text += row_text + '\n'
    return text


def has_cyrillic(text):
    return bool(re.search('[а-яА-Я]', text))


def pdf2txt(pdf_file_path):
    with open(pdf_file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ''
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()
    return text if has_cyrillic(text) else None


def rtf2txt(rtf_file_path):
    with open(rtf_file_path, 'r') as file:
        content = file.read()
    return rtf_to_txt(content)


def extract_text_from_image(image):
    text = pytesseract.image_to_string(image, lang='rus')
    return text


def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ''
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        try:
            pix = page.get_pixmap(alpha=False)
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text += extract_text_from_image(image) + '\n\n'
        except Exception as e:
            print(f"Error processing page {page_num}: {e}")
    return text


def image_to_txt(image_file):
    image = Image.open(image_file)
    text = extract_text_from_image(image)
    return text


def convert_to_txt(input_file):
    _, file_extension = os.path.splitext(input_file)
    if file_extension.lower() == '.docx':
        text = docx2txt(input_file)
        if has_image_docx(input_file):
            text += extract_images_from_docx(input_file)

    elif file_extension.lower() == '.xlsx':
        text = xlsx2txt(input_file)

    elif file_extension.lower() == '.pdf':
        text = pdf2txt(input_file)
        if has_image(input_file):
            text += extract_text_from_pdf(input_file)

    elif file_extension.lower() == '.rtf':
        text = rtf2txt(input_file)

    elif file_extension.lower() == '.doc':
        text = doc2txt(input_file)

    elif file_extension.lower() == '.txt':
        with open(input_file, 'r', encoding='utf-8') as f:
            text = f.read()
    elif file_extension.lower() in ['.jpg', '.jpeg', '.png']:
        text = image_to_txt(input_file)

    else:
        text = "Неверный формат"
    
    return text


def has_image(doc_file):
    if doc_file.lower().endswith('.pdf'):
        doc = PyPDF2.PdfReader(doc_file)
        for page_num in range(len(doc.pages)):
            page = doc.pages[page_num]
            if '/XObject' in page['/Resources']:
                return True
    elif doc_file.lower().endswith(('.doc', '.docx', '.rtf')):
        word = win32com.client.Dispatch("Word.Application")
        script_dir = os.path.dirname(os.path.abspath(__file__))
        full_path = os.path.join(script_dir, doc_file)
        doc = word.Documents.Open(full_path)
        for shape in doc.InlineShapes:
            if shape.Type == 3:
                return True
        for shape in doc.Shapes:
            if shape.Type == 13:
                return True
        doc.Close()
        word.Quit()
    return False


def has_image_docx(docx_file):
    document = Document(docx_file)
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run._element.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r':
                for child in run._element:
                    if child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing':
                        return True
    return False


def extract_images_from_docx(docx_file):
    text = ""
    document = Document(docx_file)
    for i, image in enumerate(document.inline_shapes):
        image_bytes = image._inline.graphic.graphicData.pic.blipFill.blip.embed
        for rel_id, rel in document.part.rels.items():
            if (rel.target_part.content_type == 'image/png' or rel.target_part.content_type == 'image/jpeg' or rel.target_part.content_type == 'image/bitmap' or rel.target_part.content_type == 'image/jpg') and rel_id == image_bytes:
                image_data = rel.target_part.blob
                image = Image.open(BytesIO(image_data))
                extracted_text = pytesseract.image_to_string(image, lang='rus')
                text += extracted_text + '\n\n'
                break
    return text